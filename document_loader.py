import time
import os
import win32api
import PyPDF2
from docx import Document as docxDoc
from pptx import Presentation as pptxDoc

class Document:
    def __init__(self, page_content, metadata):
        self.page_content = page_content
        self.metadata = metadata

class DocumentLoader:

    def __init__(self, whitelisted_extensions=None):
        self.whitelisted_extensions = whitelisted_extensions if whitelisted_extensions else []

    def __str__(self):
        return f"Document(\n  page_content: {self.page_content[:100]}..., \n  metadata: {self.metadata}\n)"

    def __repr__(self):
        return self.__str__()
        
    def _load_document(self, filepath):
        extension = filepath.split('.')[-1].lower()

        if extension in self.whitelisted_extensions:
            with open(filepath, 'r') as file:
                return file.read()

        if extension == 'pdf':
            with open(filepath, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                return ' '.join([reader.pages[i].extract_text() for i in range(len(reader.pages))])

        elif extension == 'docx':
            doc = docxDoc(filepath)
            return ' '.join([para.text for para in doc.paragraphs])

        elif extension == 'pptx':
            ppt = pptxDoc(filepath)
            text = []
            for slide in ppt.slides:
                for shape in slide.shapes:
                    if hasattr(shape, "text"):
                        text.append(shape.text)
            return ' '.join(text)

        else:
            raise ValueError(f"Unsupported file type: {extension}")

    def _get_file_owner(self, filepath):
        try:
            owner_SID = win32api.GetFileSecurity(filepath, win32api.OWNER_SECURITY_INFORMATION).GetSecurityDescriptorOwner()
            owner_name = win32api.LookupAccountSid(None, owner_SID)[0]
            return owner_name
        except:
            return "Unknown"

    def _convert_to_unix_timestamp(self, date_time):
        return str(int(time.mktime(time.strptime(date_time, "%a %b %d %H:%M:%S %Y"))))

    def load_from_folders(self, directories, recursive=False):
        documents_list = []
        for directory in directories:
            for root, dirs, files in os.walk(directory):
                for file in files:
                    filepath = os.path.join(root, file)
                    extension = filepath.split('.')[-1].lower()

                    if extension in ['pdf', 'docx', 'pptx'] or extension in self.whitelisted_extensions:
                        try:
                            content = self._load_document(filepath)
                        except Exception as e:
                            print(f"Error reading {filepath}: {e}")
                            content = "Error reading document."
                        updated_time = self._convert_to_unix_timestamp(time.ctime(os.path.getmtime(filepath)))
                        created_time = self._convert_to_unix_timestamp(time.ctime(os.path.getctime(filepath)))
                        author = self._get_file_owner(filepath)
                        
                        metadata = {
                            'filepath': filepath,
                            'filename': file,
                            'updated_time': updated_time,
                            'created_time': created_time,
                            'author': author or "Unknown"
                        }
                        doc = Document(page_content=content, metadata=metadata)
                        documents_list.append(doc)

                    if not recursive:
                        break
                        
        return documents_list

    def load_from_files(self, filepaths):
        documents_list = []
        for filepath in filepaths:
            extension = filepath.split('.')[-1].lower()

            if extension in ['pdf', 'docx', 'pptx'] or extension in self.whitelisted_extensions:
                content = self._load_document(filepath)
                if len(content) < 5:
                    content = "no content"
                updated_time = time.ctime(os.path.getmtime(filepath))
                created_time = time.ctime(os.path.getctime(filepath))
                author = self._get_file_owner(filepath)
                
                metadata = {
                    'filepath': filepath,
                    'filename': os.path.basename(filepath),
                    'updated_time': updated_time,
                    'created_time': created_time,
                    'author': author or "Unknown"
                }
                doc = Document(page_content=content, metadata=metadata)
                documents_list.append(doc)
                
        return documents_list

def test_document_loader():
    # Example usage:
    loader = DocumentLoader(whitelisted_extensions=["txt"])
    directories = [r"C:\Users\liamg\Scripts\AI\Langchain\pdfsearch\docs"]
    documents_list = loader.load_from_folders(directories, recursive=True)

    for doc in documents_list:
        print("-" * 50)  # print a separator for clarity between documents
        print(f"Page Content: '{doc.page_content[:1000]}'")  # For brevity, only the first 1000 characters
        for key, value in doc.metadata.items():
            print(f"{key}: '{value}'")

if __name__ == '__main__':
    test_document_loader()